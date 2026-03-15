"""DOCX 파일 파싱 모듈

학습 문서 특성: 본문 대부분이 Table 셀 안에 존재.
Paragraph(제목 등) + Table 셀 텍스트를 모두 추출하여 raw_text에 합산.
"""
import re
from pathlib import Path

from docx import Document
from pydantic import BaseModel


class Section(BaseModel):
    """문서 섹션"""
    heading: str = ""
    level: int = 0
    content: str = ""
    items: list[str] = []


class TableData(BaseModel):
    """테이블 데이터"""
    headers: list[str] = []
    rows: list[list[str]] = []


class ParsedDocument(BaseModel):
    """파싱된 문서"""
    file_id: str
    filename: str
    domain: str
    format: str
    sections: list[Section] = []
    tables: list[TableData] = []
    raw_text: str = ""
    metadata: dict[str, str] = {}


def extract_file_id(filename: str) -> tuple[str, str]:
    """파일명에서 도메인과 ID 추출.

    패턴: {도메인}_{번호}_{제목}.docx
    예: CA_001_멧칼프 법칙.docx → ("CA", "CA_001")
    """
    stem = Path(filename).stem
    match = re.match(r"^([A-Z]{2})_(\d+(?:\.\d+)?)", stem)
    if match:
        domain = match.group(1)
        file_id = f"{domain}_{match.group(2)}"
        return domain, file_id
    return "", stem


def _extract_metadata_from_tables(tables: list[TableData]) -> dict[str, str]:
    """첫 번째 메타 테이블에서 토픽명, 분류, 키워드 등 추출.

    학습 문서 표준 포맷:
      | 토픽 이름 (중) | 멧칼프의 법칙 |
      | 분류          | CA > 법칙 일반  |
      | 키워드(암기)   | 유용성, 임계값  |
    """
    meta: dict[str, str] = {}
    if not tables:
        return meta

    first = tables[0]
    # 첫 테이블의 header + rows를 key-value로 시도
    all_rows = [first.headers] + first.rows
    for row in all_rows:
        if len(row) >= 2:
            key = row[0].strip()
            val = row[1].strip()
            if key and val:
                # 주요 메타데이터 키 정규화
                key_lower = key.replace(" ", "")
                if "토픽" in key_lower or "이름" in key_lower:
                    meta["topic_name"] = val
                elif "분류" in key_lower:
                    meta["classification"] = val
                elif "키워드" in key_lower or "암기" in key_lower:
                    meta["keywords"] = val
                elif "암기법" in key_lower:
                    meta["mnemonic"] = val
    return meta


def parse_docx(filepath: Path) -> ParsedDocument:
    """DOCX 파일을 파싱하여 구조화된 데이터 반환.

    Paragraph + Table 셀 텍스트를 모두 추출.
    """
    doc = Document(str(filepath))
    domain, file_id = extract_file_id(filepath.name)

    sections: list[Section] = []
    tables: list[TableData] = []
    raw_lines: list[str] = []
    current_section = Section()

    # ── 1. Paragraph 파싱 ──
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        raw_lines.append(text)

        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            if current_section.heading or current_section.content:
                sections.append(current_section)
            level = int(style_name.replace("Heading ", "").replace("Heading", "1") or "1")
            current_section = Section(heading=text, level=level)
        elif re.match(r"^[가-힣]\.", text):
            current_section.items.append(text)
            current_section.content += text + "\n"
        elif re.match(r"^[IVX]+\.", text) or re.match(r"^\d+\.", text):
            # I. II. III. 또는 1. 2. 3. 형태의 대제목
            if current_section.heading or current_section.content:
                sections.append(current_section)
            current_section = Section(heading=text, level=1)
        else:
            current_section.content += text + "\n"

    if current_section.heading or current_section.content:
        sections.append(current_section)

    # ── 2. Table 파싱 (본문 핵심) ──
    seen_cells: set[str] = set()  # 병합 셀 중복 제거

    for table in doc.tables:
        rows_data: list[list[str]] = []
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                # 병합 셀 중복 제거 (docx에서 병합 셀은 같은 텍스트를 여러 번 반환)
                cell_key = f"{id(table)}_{cell._element.xml[:50]}_{cell_text[:30]}"
                if cell_key not in seen_cells:
                    seen_cells.add(cell_key)
                    cells.append(cell_text)
                    # 테이블 셀 텍스트도 raw_text에 포함
                    if cell_text and len(cell_text) > 3:
                        raw_lines.append(cell_text)
            if cells:
                rows_data.append(cells)

        if rows_data:
            td = TableData(
                headers=rows_data[0],
                rows=rows_data[1:] if len(rows_data) > 1 else [],
            )
            tables.append(td)

    # ── 3. 테이블 내 섹션 구조 추출 ──
    # 대형 콘텐츠 테이블에서 추가 섹션 추출
    for table_data in tables:
        for row in table_data.rows:
            for cell_text in row:
                if not cell_text:
                    continue
                for line in cell_text.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    # 로마숫자 대제목 (I. II. III.)
                    if re.match(r"^[IVX]+\.\s", line):
                        if current_section.heading or current_section.content:
                            sections.append(current_section)
                        current_section = Section(heading=line, level=1)
                    elif re.match(r"^[가-힣]\.\s", line):
                        current_section.items.append(line)
                        current_section.content += line + "\n"
                    elif re.match(r"^\d+\)\s", line) or re.match(r"^-\s", line):
                        current_section.content += line + "\n"

    if current_section.heading or current_section.content:
        # 이미 추가된 마지막 섹션과 중복 방지
        if not sections or sections[-1].heading != current_section.heading:
            sections.append(current_section)

    # ── 4. 메타데이터 추출 ──
    metadata = _extract_metadata_from_tables(tables)

    return ParsedDocument(
        file_id=file_id,
        filename=filepath.name,
        domain=domain,
        format="docx",
        sections=sections,
        tables=tables,
        raw_text="\n".join(raw_lines),
        metadata=metadata,
    )
